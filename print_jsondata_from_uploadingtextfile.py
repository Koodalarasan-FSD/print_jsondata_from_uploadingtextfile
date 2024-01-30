from flask import Flask, render_template, request, send_file
from docx import Document
import os
import json

app = Flask(__name__)

@app.route('/')
def index():
    return """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta http-equiv="X-UA-Compatible" content="IE=edge">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>JSON to Word</title>
        </head>
        <body>
            <h1>JSON to Word Converter</h1>
            <form action="/generate_word" method="post" enctype="multipart/form-data">
                <label for="jsonFile">Upload JSON Text File (.txt):</label>
                <input type="file" id="jsonFile" name="jsonFile" accept=".txt" required><br>
                <button type="submit">Generate Word</button>
            </form>
        </body>
        </html>



    """

@app.route('/generate_word', methods=['POST'])
def generate_word():
    # Check if the file is present in the request
    if 'jsonFile' not in request.files:
        return "No file provided"

    file = request.files['jsonFile']

    # Check if the file has a .txt extension
    if not file.filename.endswith('.txt'):
        return "Invalid file format. Please upload a text file with a .txt extension."

    try:
        # Read the JSON data from the file
        json_data = json.load(file)
    except json.JSONDecodeError as e:
        return f"Error parsing JSON: {str(e)}"

    # Create a new Word document
    doc = Document()

    # Add content to the Word document
    doc.add_heading('JSON Data', level=1)
    for key, value in json_data.items():
        doc.add_paragraph(f"{key}: {value}")

    # Save the Word document in the current working directory
    save_path = os.path.join(os.getcwd(), 'output.docx')
    doc.save(save_path)

    # Send the Word document as a response
    return send_file(save_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
